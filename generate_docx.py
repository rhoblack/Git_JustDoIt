#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Git 입문 교재 Word 문서 생성 스크립트
HTML 원본을 읽어 Microsoft Word 형식의 출판용 교재를 생성합니다.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os
from html.parser import HTMLParser

# 색상 정의
CHAPTER_COLORS = {
    1: (102, 126, 234),   # 보라
    2: (33, 150, 243),    # 파랑
    3: (255, 152, 0),     # 주황
    4: (244, 67, 54),     # 빨강
    5: (156, 39, 176),    # 자주
}

class HTMLToDocxParser(HTMLParser):
    """HTML을 Word 문서로 변환"""

    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.current_para = None
        self.in_code = False
        self.in_strong = False
        self.in_em = False

    def handle_starttag(self, tag, attrs):
        if tag == "h2":
            self.current_para = self.doc.add_heading("", level=2)
        elif tag == "h3":
            self.current_para = self.doc.add_heading("", level=3)
        elif tag == "p":
            self.current_para = self.doc.add_paragraph()
        elif tag == "strong":
            self.in_strong = True
        elif tag == "em":
            self.in_em = True
        elif tag == "code" or tag == "pre":
            self.in_code = True
        elif tag == "ul":
            pass  # 리스트 처리는 간단히
        elif tag == "li":
            self.current_para = self.doc.add_paragraph("", style="List Bullet")

    def handle_endtag(self, tag):
        if tag in ["h2", "h3", "p", "li", "ul"]:
            self.current_para = None
        elif tag == "strong":
            self.in_strong = False
        elif tag == "em":
            self.in_em = False
        elif tag in ["code", "pre"]:
            self.in_code = False

    def handle_data(self, data):
        text = data.strip()
        if not text:
            return

        if self.current_para is None:
            self.current_para = self.doc.add_paragraph()

        run = self.current_para.add_run(text)

        if self.in_strong:
            run.bold = True
        if self.in_em:
            run.italic = True
        if self.in_code:
            run.font.name = 'Courier New'
            run.font.color.rgb = RGBColor(44, 44, 44)

def extract_html_content(html_file):
    """HTML 파일에서 섹션과 내용 추출"""

    with open(html_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # 헤더 추출 (장 제목)
    header_match = re.search(r'<header[^>]*>.*?<h1[^>]*>([^<]+)</h1>', content, re.DOTALL)
    chapter_title = header_match.group(1) if header_match else ""

    # 섹션별 내용 추출
    sections = re.findall(r'<section[^>]*>(.*?)</section>', content, re.DOTALL)

    extracted = []
    for section in sections:
        # h2 제목
        title_match = re.search(r'<h2[^>]*>([^<]+)</h2>', section)
        title = title_match.group(1) if title_match else ""

        # 단락들
        paragraphs = re.findall(r'<p[^>]*>([^<]+)</p>', section)

        # 코드 블록
        codes = re.findall(r'<code[^>]*>([^<]+)</code>', section)

        if title:
            extracted.append({
                "title": title.strip(),
                "content": [p.strip() for p in paragraphs if p.strip()],
                "code": [c.strip() for c in codes if c.strip()]
            })

    return chapter_title, extracted

def set_cell_background(cell, fill_color):
    """테이블 셀 배경색 설정"""
    cell_xml = cell._element
    cell_pr = cell_xml.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), fill_color)
    cell_pr.append(shade_obj)

def create_cover_page(doc):
    """표지 생성"""
    # 제목
    title = doc.add_paragraph()
    title_run = title.add_run("Git 입문 교재")
    title_run.font.size = Pt(48)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(102, 126, 234)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(12)

    # 부제
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("완전 가이드")
    subtitle_run.font.size = Pt(32)
    subtitle_run.font.color.rgb = RGBColor(150, 150, 150)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.space_after = Pt(48)

    # 설명
    desc = doc.add_paragraph("Git을 처음 배우는 개발자를 위한\n단계별 학습 가이드")
    desc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_run = desc.runs[0]
    desc_run.font.size = Pt(18)
    desc_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

def create_table_of_contents(doc):
    """목차 생성"""
    title = doc.add_heading("목차", level=1)

    chapters = [
        "1장 Git 시작하기",
        "2장 기본 작업",
        "3장 협업하기",
        "4장 문제 해결",
        "5장 실전 상황",
    ]

    for chapter in chapters:
        p = doc.add_paragraph(chapter, style="List Number")
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_page_break()

def add_chapter_to_doc(doc, chapter_num, chapter_title, sections):
    """챕터를 Word 문서에 추가"""

    # 챕터 제목
    heading = doc.add_heading(f"{chapter_num}장 {chapter_title}", level=1)
    heading.runs[0].font.color.rgb = RGBColor(*CHAPTER_COLORS[chapter_num])

    # 각 섹션 추가
    for section in sections:
        # 섹션 제목
        sub_heading = doc.add_heading(section["title"], level=2)
        sub_heading.runs[0].font.color.rgb = RGBColor(*CHAPTER_COLORS[chapter_num])

        # 내용
        for para_text in section["content"]:
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.paragraph_format.line_spacing = 1.5

        # 코드 블록
        if section["code"]:
            # 코드 배경 테이블로 표현
            table = doc.add_table(rows=1, cols=1)
            table.autofit = False
            cell = table.rows[0].cells[0]

            # 배경색 설정 (어두운 회색)
            set_cell_background(cell, "2a2a2a")

            # 코드 텍스트
            para = cell.paragraphs[0]
            for code_line in section["code"]:
                run = para.add_run(f"$ {code_line}\n")
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)

            # 테이블 스타일
            tbl = table._element
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)

    doc.add_page_break()

def generate_docx():
    """Word 문서 생성 메인 함수"""

    print("[*] Git 입문 교재 Word 문서 생성 시작...")

    # 문서 생성
    doc = Document()

    # 기본 스타일 설정
    style = doc.styles['Normal']
    style.font.name = 'Segoe UI'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # 표지
    create_cover_page(doc)

    # 목차
    create_table_of_contents(doc)

    # 각 장 추가
    chapters = [
        (1, "Git 시작하기"),
        (2, "기본 작업"),
        (3, "협업하기"),
        (4, "문제 해결"),
        (5, "실전 상황"),
    ]

    for chapter_num, chapter_title in chapters:
        html_file = f"chapter{chapter_num}.html"

        if not os.path.exists(html_file):
            print(f"   [!] {html_file} 파일을 찾을 수 없습니다.")
            continue

        print(f"[*] Chapter {chapter_num}: {chapter_title} 처리 중...")

        title, sections = extract_html_content(html_file)
        add_chapter_to_doc(doc, chapter_num, chapter_title, sections)

    # 부록: 자주 묻는 질문
    doc.add_heading("부록: 자주 묻는 질문 (FAQ)", level=1)

    faqs = [
        ("Git을 배우려면 프로그래밍을 알아야 하나요?", "아니요. 본 교재는 비프로그래머도 이해할 수 있도록 작성되었습니다."),
        ("Git과 GitHub의 차이가 뭔가요?", "Git은 버전 관리 도구이고, GitHub는 Git을 사용하는 원격 저장소 서비스입니다."),
        ("CLI와 GUI 중 어느 것을 배워야 하나요?", "본 교재는 CLI(명령어)를 기본으로 하고 있으며, GUI 도구도 함께 소개합니다."),
        ("브랜치는 언제 필요한가요?", "여러 사람이 함께 작업하거나, 여러 기능을 동시에 개발할 때 필요합니다."),
        ("Merge conflict는 왜 발생하나요?", "같은 파일의 같은 부분을 두 명 이상이 수정했을 때 발생합니다."),
    ]

    for q, a in faqs:
        p = doc.add_heading(q, level=3)
        doc.add_paragraph(a)

    doc.add_page_break()

    # 명령어 참고표
    doc.add_heading("명령어 참고표", level=1)

    commands = [
        ("git init", "새 저장소 초기화"),
        ("git add", "변경사항 스테이징"),
        ("git commit", "커밋 생성"),
        ("git push", "원격 저장소에 업로드"),
        ("git pull", "원격 저장소에서 가져오기"),
        ("git branch", "브랜치 관리"),
        ("git merge", "브랜치 병합"),
        ("git reset", "HEAD 초기화"),
        ("git revert", "이전 커밋으로 되돌리기"),
        ("git stash", "변경사항 임시 저장"),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'

    # 헤더행
    header_cells = table.rows[0].cells
    header_cells[0].text = "명령어"
    header_cells[1].text = "설명"

    # 명령어들
    for cmd, desc in commands:
        row_cells = table.add_row().cells
        row_cells[0].text = cmd
        row_cells[1].text = desc

    # 저장
    output_dir = "output/docx"
    os.makedirs(output_dir, exist_ok=True)
    output_file = f"{output_dir}/Git_입문_교재.docx"

    doc.save(output_file)

    print(f"[OK] Word 문서 생성 완료: {output_file}")
    print(f"[*] 총 페이지: 약 {len(doc.element.body)} 섹션")

if __name__ == "__main__":
    generate_docx()
